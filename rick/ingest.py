"""
Ingestione documenti v11 — Robust, con retry, progresso e statistiche.
Supporta: TXT, MD, PDF, DOCX, HTML, CSV, JSON, ZIP (anche annidati).
"""
import argparse
import logging
import time
import zipfile
import tempfile
import shutil
import os
from pathlib import Path
from typing import List, Tuple
from rick.memory import add_knowledge

# ── Dipendenze opzionali ─────────────────────────────────────────────────────
try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

try:
    import chardet
except ImportError:
    chardet = None

logging.basicConfig(level=logging.INFO, format='%(asctime)s [%(levelname)s] %(message)s')
logger = logging.getLogger(__name__)

CHUNK_SIZE = 400
SUPPORTED_EXTS = {".txt", ".md", ".pdf", ".docx", ".html", ".htm", ".py", ".sh", ".js", ".json", ".csv", ".xml", ".rst", ".yaml", ".yml"}
MAX_CHUNKS_PER_FILE = 100  # sicurezza: evita ingestion di file enormi
RETRY_DELAY = 2  # secondi tra tentativi


def detect_encoding(filepath: Path) -> str:
    """Rileva la codifica di un file, con fallback a UTF-8."""
    if chardet:
        try:
            with open(filepath, 'rb') as f:
                raw = f.read(10000)
            result = chardet.detect(raw)
            return result.get('encoding', 'utf-8') or 'utf-8'
        except:
            return 'utf-8'
    return 'utf-8'


def chunk_text(text: str, size: int = CHUNK_SIZE) -> List[str]:
    """Divide il testo in chunk di ~size parole, preservando i paragrafi quando possibile."""
    paragraphs = text.split('\n')
    chunks = []
    current = []
    current_len = 0
    
    for para in paragraphs:
        words = para.split()
        if not words:
            if current:
                chunks.append(' '.join(current))
                current, current_len = [], 0
            continue
        if current_len + len(words) <= size:
            current.extend(words)
            current_len += len(words)
        else:
            if current:
                chunks.append(' '.join(current))
            current = words.copy()
            current_len = len(words)
    if current:
        chunks.append(' '.join(current))
    return chunks


def read_file(filepath: Path) -> str | None:
    """Legge il contenuto testuale da vari formati."""
    try:
        suffix = filepath.suffix.lower()
        
        if suffix == ".pdf":
            if not PdfReader:
                logger.warning("[ingest] pypdf non installato, salto PDF")
                return None
            reader = PdfReader(filepath)
            return "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
        
        elif suffix == ".docx":
            if not Document:
                logger.warning("[ingest] python-docx non installato, salto DOCX")
                return None
            doc = Document(filepath)
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            return "\n".join(parts)
        
        elif suffix in [".html", ".htm"]:
            if not BeautifulSoup:
                logger.warning("[ingest] beautifulsoup4 non installato, salto HTML")
                return None
            content = filepath.read_text(encoding="utf-8", errors="ignore")
            soup = BeautifulSoup(content, "html.parser")
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            text = soup.get_text(separator="\n")
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            return "\n".join(lines)
        
        else:
            # Testo semplice
            encodings = [detect_encoding(filepath), 'utf-8', 'latin-1', 'cp1252']
            for enc in sorted(set(encodings), key=lambda x: encodings.index(x)):
                try:
                    return filepath.read_text(encoding=enc)
                except:
                    continue
            return filepath.read_text(encoding="utf-8", errors="ignore")
    
    except Exception as e:
        logger.error(f"[ingest] Errore lettura {filepath.name}: {e}")
        return None


def ingest_file(filepath: Path) -> Tuple[int, int]:
    """Ingerisce un singolo file. Restituisce (chunks_added, chunks_total)."""
    if filepath.suffix.lower() not in SUPPORTED_EXTS:
        return 0, 0
    
    logger.info(f"[ingest] 📄 {filepath.name}")
    text = read_file(filepath)
    if not text or len(text.strip()) < 10:
        return 0, 0
    
    chunks = chunk_text(text)
    if len(chunks) > MAX_CHUNKS_PER_FILE:
        logger.warning(f"[ingest] {filepath.name} ha {len(chunks)} chunk, limito a {MAX_CHUNKS_PER_FILE}")
        chunks = chunks[:MAX_CHUNKS_PER_FILE]
    
    added = 0
    for i, chunk in enumerate(chunks):
        retries = 0
        while retries < 3:
            try:
                add_knowledge(chunk, source_name=f"{filepath.name}#chunk{i}")
                added += 1
                break
            except Exception as e:
                retries += 1
                logger.warning(f"[ingest] Chunk {i} fallito (tentativo {retries}/3): {e}")
                time.sleep(RETRY_DELAY * retries)
        if (i + 1) % 20 == 0:
            logger.info(f"[ingest]   {i+1}/{len(chunks)} chunk elaborati...")
    
    return added, len(chunks)


def _process_directory(directory: Path) -> Tuple[int, int, int]:
    """Elabora ricorsivamente una directory."""
    logger.info(f"[ingest] 📁 Scansione directory: {directory.name}")
    all_files = []
    for ext in SUPPORTED_EXTS:
        all_files.extend(list(directory.rglob(f"*{ext}")))
    all_files = sorted(set(all_files))
    logger.info(f"[ingest] Trovati {len(all_files)} file supportati")
    
    total_added, total_chunks = 0, 0
    for f in all_files:
        added, chunks = ingest_file(f)
        total_added += added
        total_chunks += chunks
    
    return len(all_files), total_added, total_chunks


def ingest_anything(path_str: str) -> dict:
    """Funzione principale. Gestisce file, directory e ZIP."""
    path = Path(path_str).expanduser().resolve()
    if not path.exists():
        logger.error(f"[ingest] Percorso non trovato: {path}")
        return {}

    if path.suffix.lower() == ".zip":
        logger.info(f"[ingest] 📦 Estrazione ZIP: {path.name}")
        with tempfile.TemporaryDirectory() as tmpdir:
            try:
                with zipfile.ZipFile(path, 'r') as zip_ref:
                    zip_ref.extractall(tmpdir)
                for item in Path(tmpdir).rglob("*.zip"):
                    logger.info(f"[ingest]   → ZIP annidato: {item.name}")
                    with tempfile.TemporaryDirectory() as inner_tmp:
                        try:
                            with zipfile.ZipFile(item, 'r') as inner_zip:
                                inner_zip.extractall(inner_tmp)
                            shutil.copytree(inner_tmp, tmpdir, dirs_exist_ok=True)
                        except: pass
                return ingest_anything(tmpdir)
            except Exception as e:
                logger.error(f"[ingest] Errore ZIP: {e}")
                return {}

    if path.is_dir():
        files, added, total = _process_directory(path)
        logger.info(f"[ingest] ✅ Completato: {files} file, {added}/{total} chunk aggiunti")
        return {"files": files, "chunks_added": added, "chunks_total": total}

    if path.is_file():
        added, total = ingest_file(path)
        logger.info(f"[ingest] ✅ Completato: {added}/{total} chunk aggiunti")
        return {"files": 1, "chunks_added": added, "chunks_total": total}

    return {}


def main():
    parser = argparse.ArgumentParser(description="Rick Knowledge Ingestor")
    parser.add_argument("path", type=str, help="File, directory o ZIP")
    args = parser.parse_args()
    
    t0 = time.time()
    stats = ingest_anything(args.path)
    if stats:
        elapsed = time.time() - t0
        logger.info(f"[ingest] Tempo totale: {elapsed:.1f}s")


if __name__ == "__main__":
    main()
